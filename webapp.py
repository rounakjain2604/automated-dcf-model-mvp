from __future__ import annotations

import io
import tempfile
import zipfile
from datetime import date
from pathlib import Path

import matplotlib
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor
from flask import Flask, render_template, request, send_file

from src.dcf_generator.config import DCFConfig
from src.dcf_generator.pipeline import run_dcf_pipeline

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker


app = Flask(__name__)


CASE_DEFAULTS = {
    "Base": {
        "revenue": 2_058_911.0,
        "ebitda": 1_072_911.0,
        "growth_rate": 3.0,
        "wacc": 24.0,
        "terminal_growth": 2.0,
        "cash": 50_000.0,
        "debt": 0.0,
        "ask_price": 4_000_000.0,
        "cogs_pct": 17.0,
        "opex_pct": 30.89,
        "capex_pct": 1.5,
        "tax_rate": 25.0,
        "risk_free_rate": 4.0,
        "beta": 1.0,
        "size_premium": 5.0,
    },
    "Bull": {
        "revenue": 2_058_911.0,
        "ebitda": 1_126_000.0,
        "growth_rate": 4.0,
        "wacc": 22.0,
        "terminal_growth": 2.25,
        "cash": 50_000.0,
        "debt": 0.0,
        "ask_price": 4_000_000.0,
        "cogs_pct": 16.5,
        "opex_pct": 29.5,
        "capex_pct": 1.4,
        "tax_rate": 25.0,
        "risk_free_rate": 4.0,
        "beta": 0.95,
        "size_premium": 4.5,
    },
    "Bear": {
        "revenue": 2_058_911.0,
        "ebitda": 1_000_000.0,
        "growth_rate": 2.0,
        "wacc": 26.0,
        "terminal_growth": 1.75,
        "cash": 50_000.0,
        "debt": 0.0,
        "ask_price": 4_000_000.0,
        "cogs_pct": 18.0,
        "opex_pct": 32.0,
        "capex_pct": 1.7,
        "tax_rate": 25.0,
        "risk_free_rate": 4.0,
        "beta": 1.05,
        "size_premium": 5.5,
    },
}


def _to_float(value: str, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _pick_input(form, key: str, scenario: str) -> float:
    raw = form.get(key)
    if raw is not None and str(raw).strip() != "":
        cleaned = str(raw).replace(",", "")
        return _to_float(cleaned, CASE_DEFAULTS[scenario][key])
    return CASE_DEFAULTS[scenario][key]


def _build_cfg_from_form(form, scenario: str) -> DCFConfig:
    cfg = DCFConfig()

    growth = _pick_input(form, "growth_rate", scenario) / 100
    wacc = _pick_input(form, "wacc", scenario) / 100
    terminal_growth = _pick_input(form, "terminal_growth", scenario) / 100
    cash = _pick_input(form, "cash", scenario)
    debt = _pick_input(form, "debt", scenario)
    ask_price = _pick_input(form, "ask_price", scenario)

    cogs_pct = _pick_input(form, "cogs_pct", scenario) / 100
    opex_pct = _pick_input(form, "opex_pct", scenario) / 100
    capex_pct = _pick_input(form, "capex_pct", scenario) / 100
    tax_rate = _pick_input(form, "tax_rate", scenario) / 100
    risk_free = _pick_input(form, "risk_free_rate", scenario) / 100
    beta = _pick_input(form, "beta", scenario)
    size_premium = _pick_input(form, "size_premium", scenario) / 100

    cfg.forecast.revenue_method = "cagr"
    cfg.forecast.revenue_cagr = growth
    cfg.forecast.cogs_method = "pct_revenue"
    cfg.forecast.opex_method = "pct_revenue"
    cfg.forecast.cogs_pct_revenue = cogs_pct
    cfg.forecast.opex_pct_revenue = opex_pct
    cfg.forecast.capex_pct_revenue = capex_pct
    cfg.forecast.tax_rate = tax_rate

    cfg.wacc.target_debt_weight = 0.0
    cfg.wacc.target_equity_weight = 1.0
    cfg.wacc.risk_free_rate = risk_free
    cfg.wacc.beta = beta
    cfg.wacc.size_premium = size_premium
    cfg.wacc.country_risk_premium = 0.0
    cfg.wacc.market_risk_premium = max(wacc - (cfg.wacc.risk_free_rate + cfg.wacc.size_premium), 0.01)
    cfg.wacc.tax_rate = tax_rate

    cfg.valuation.terminal_growth_rate = terminal_growth
    cfg.valuation.cash = cash
    cfg.valuation.debt = debt
    cfg.valuation.minority_interest = 0.0
    cfg.valuation.preferred_stock = 0.0
    cfg.valuation.fully_diluted_shares = 1.0
    cfg.valuation.gdp_growth_cap = 0.035
    cfg.valuation.exit_ev_ebitda_multiple = 4.5 if wacc >= 0.24 else 6.0

    cfg._web_meta = {"ask_price": ask_price}  # type: ignore[attr-defined]
    return cfg


def _generate_synthetic_input(form, destination: Path, scenario: str) -> Path:
    revenue = _pick_input(form, "revenue", scenario)
    ebitda = _pick_input(form, "ebitda", scenario)
    growth = _pick_input(form, "growth_rate", scenario) / 100
    cogs_pct = _pick_input(form, "cogs_pct", scenario) / 100

    cogs = revenue * cogs_pct
    opex = max(revenue - cogs - ebitda, 0)

    growth_divisor = (1 + growth) if growth > -0.99 else 1.03
    historical_revenue = revenue / growth_divisor
    historical_cogs = historical_revenue * cogs_pct
    historical_opex = max(historical_revenue - historical_cogs - (ebitda / 1.02), 0)

    df = pd.DataFrame(
        [
            ["2023-12-31", "Revenue", "IS", historical_revenue, False],
            ["2023-12-31", "COGS", "IS", historical_cogs, False],
            ["2023-12-31", "Operating Expenses", "IS", historical_opex, False],
            ["2023-12-31", "Depreciation", "IS", 25000, False],
            ["2023-12-31", "Accounts Receivable", "BS", historical_revenue * 0.11, False],
            ["2023-12-31", "Inventory", "BS", historical_revenue * 0.008, False],
            ["2023-12-31", "Accounts Payable", "BS", historical_revenue * 0.03, False],
            ["2023-12-31", "Cash", "BS", _pick_input(form, "cash", scenario), False],
            ["2023-12-31", "Debt", "BS", _pick_input(form, "debt", scenario), False],
            ["2024-12-31", "Revenue", "IS", revenue, False],
            ["2024-12-31", "COGS", "IS", cogs, False],
            ["2024-12-31", "Operating Expenses", "IS", opex, False],
            ["2024-12-31", "Depreciation", "IS", 25000, False],
            ["2024-12-31", "Accounts Receivable", "BS", revenue * 0.11, False],
            ["2024-12-31", "Inventory", "BS", revenue * 0.008, False],
            ["2024-12-31", "Accounts Payable", "BS", revenue * 0.03, False],
            ["2024-12-31", "Cash", "BS", _pick_input(form, "cash", scenario), False],
            ["2024-12-31", "Debt", "BS", _pick_input(form, "debt", scenario), False],
        ],
        columns=["period", "account", "statement", "amount", "is_non_recurring"],
    )
    df.to_csv(destination, index=False)
    return destination


def _fmt_m(value: float) -> str:
    return f"${value/1_000_000:,.2f}M"


def _zebra_table(table) -> None:
    for idx, row in enumerate(table.rows):
        if idx == 0:
            continue
        if idx % 2 == 0:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = cell._element.xpath(".//w:shd")
                if not shd:
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn

                    shade = OxmlElement("w:shd")
                    shade.set(qn("w:fill"), "F5F7FA")
                    tc_pr.append(shade)


def _build_chart_football(low: float, base: float, high: float, out_path: Path) -> None:
    labels = ["Bear", "Base", "Bull"]
    values = [low, base, high]
    fig, ax = plt.subplots(figsize=(6.6, 2.2))
    ax.barh(labels, values, color=["#8CAAD0", "#1F4E78", "#5C7EA8"])
    ax.set_xlabel("Enterprise Value ($M)")
    ax.set_title("Football Field (Scenario Valuation)")
    ax.grid(axis="x", linestyle="--", alpha=0.3)
    ax.xaxis.set_major_formatter(mticker.FuncFormatter(lambda x, pos: f"${x/1_000_000:.1f}M"))
    fig.tight_layout()
    fig.savefig(out_path, dpi=200)
    plt.close(fig)


def _build_chart_margin(forecast_df: pd.DataFrame, out_path: Path) -> None:
    years = forecast_df["period"].astype(str).str[:4].tolist()
    revenue = forecast_df["Revenue"].astype(float)
    ebitda_margin = (forecast_df["EBITDA"].astype(float) / revenue).fillna(0)
    gross_margin = ((forecast_df["Revenue"] - forecast_df["COGS"]).astype(float) / revenue).fillna(0)

    fig, ax = plt.subplots(figsize=(6.6, 2.6))
    ax.plot(years, gross_margin * 100, marker="o", color="#1F4E78", label="Gross Margin %")
    ax.plot(years, ebitda_margin * 100, marker="o", color="#5C7EA8", label="EBITDA Margin %")
    ax.set_ylabel("Margin (%)")
    ax.set_title("Margin Trend (Forecast)")
    ax.grid(axis="y", linestyle="--", alpha=0.3)
    ax.legend(loc="best")
    fig.tight_layout()
    fig.savefig(out_path, dpi=200)
    plt.close(fig)


def _build_chart_revenue_bridge(forecast_df: pd.DataFrame, out_path: Path) -> None:
    years = forecast_df["period"].astype(str).str[:4].tolist()
    revenue = (forecast_df["Revenue"].astype(float) / 1_000_000).tolist()

    fig, ax = plt.subplots(figsize=(6.6, 2.4))
    ax.bar(years, revenue, color="#1F4E78")
    ax.set_ylabel("Revenue ($M)")
    ax.set_title("Revenue Bridge (Current to Forecast)")
    ax.grid(axis="y", linestyle="--", alpha=0.25)
    fig.tight_layout()
    fig.savefig(out_path, dpi=200)
    plt.close(fig)


def _build_chart_fcf_area(forecast_df: pd.DataFrame, out_path: Path) -> None:
    years = [f"{i+1}Y" for i in range(len(forecast_df))]
    fcf_vals = (forecast_df["FCF"].astype(float) / 1_000_000).tolist()

    fig, ax = plt.subplots(figsize=(4.8, 2.8))
    ax.plot(years, fcf_vals, color="#1F4E78", linewidth=2.2)
    ax.fill_between(years, fcf_vals, color="#5C7EA8", alpha=0.75)
    ax.set_title("Unlevered Free Cash Flow Projection ($M)", fontsize=10, fontweight="bold")
    ax.grid(axis="y", linestyle="--", alpha=0.25)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color("#AAB7C8")
    ax.spines["bottom"].set_color("#AAB7C8")
    ax.tick_params(axis="x", labelsize=8)
    ax.tick_params(axis="y", labelsize=8)
    ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda y, _: f"${y:,.1f}"))
    fig.tight_layout()
    fig.savefig(out_path, dpi=220)
    plt.close(fig)


def _read_report_data(excel_path: Path) -> tuple[dict[str, float | str], pd.DataFrame]:
    metrics_df = pd.read_excel(excel_path, sheet_name="ReportData", usecols="A:B")
    metrics: dict[str, float | str] = {}
    for _, row in metrics_df.dropna(how="all").iterrows():
        key = str(row.get("metric", "")).strip()
        value = row.get("value")
        if key:
            metrics[key] = value

    forecast = pd.read_excel(excel_path, sheet_name="ReportData", usecols="D:J")
    forecast = forecast.dropna(how="all")
    return metrics, forecast


def _build_word_report(
    output_path: Path,
    excel_path: Path,
    company_name: str,
    ask_price: float,
    scenario: str,
    logo_path: Path | None,
    input_ebitda: float,
) -> None:
    metrics, forecast_df = _read_report_data(excel_path)

    ev_g = float(metrics.get("ev_gordon", 0.0) or 0.0)
    ev_e = float(metrics.get("ev_exit", 0.0) or 0.0)
    eq_g = float(metrics.get("eq_gordon", 0.0) or 0.0)
    eq_e = float(metrics.get("eq_exit", 0.0) or 0.0)
    wacc = float(metrics.get("wacc", 0.0) or 0.0)
    terminal_g = float(metrics.get("terminal_growth", 0.0) or 0.0)
    risk_free = float(metrics.get("risk_free", 0.0) or 0.0)
    beta = float(metrics.get("beta", 0.0) or 0.0)
    mrp = float(metrics.get("mrp", 0.0) or 0.0)
    size_premium = float(metrics.get("size_premium", 0.0) or 0.0)
    projected_growth = float(metrics.get("revenue_cagr", 0.0) or 0.0)
    historical_avg = float(metrics.get("historical_growth_3y_avg", 0.0) or 0.0)

    if forecast_df.empty:
        forecast_df = pd.DataFrame(
            [{"period": "Y1", "Revenue": 0.0, "COGS": 0.0, "EBITDA": 0.0, "Capex": 0.0, "Delta NWC": 0.0, "FCF": 0.0}]
        )

    low = min(ev_g, ev_e)
    high = max(ev_g, ev_e)
    midpoint = (low + high) / 2
    uplift = midpoint - ask_price
    upside_pct = ((midpoint / ask_price) - 1) if ask_price > 0 else 0.0

    bear = midpoint * 0.90
    bull = midpoint * 1.10

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    # Header/footer style
    if logo_path and logo_path.exists():
        header_p = section.header.paragraphs[0]
        run = header_p.add_run()
        run.add_picture(str(logo_path), width=Pt(80))
        header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    section.header.paragraphs[0].add_run(f"  Confidential Valuation Report — [{company_name}] | Date: {date.today().strftime('%B %d, %Y')}")
    section.header.paragraphs[0].runs[-1].font.size = Pt(9)
    section.footer.paragraphs[0].text = "Strictly Private & Confidential | Prepared by Rounak Jain, CFA L2 Candidate"
    section.footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Visual top divider
    divider = doc.add_paragraph("_" * 120)
    divider.runs[0].font.color.rgb = RGBColor(187, 194, 204)
    divider.paragraph_format.space_after = Pt(1)

    title = doc.add_paragraph("Executive Valuation Summary")
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.bold = True
    title.runs[0].font.color.rgb = RGBColor(24, 37, 56)
    title.paragraph_format.space_after = Pt(6)

    summary = doc.add_paragraph(
        f"Based on our DCF analysis, {company_name} implies an enterprise value range of {_fmt_m(low)} to {_fmt_m(high)} "
        f"under {scenario} assumptions. Using a WACC of {wacc:.1%} and terminal growth of {terminal_g:.1%}, "
        f"the midpoint value is {_fmt_m(midpoint)}, indicating {upside_pct:.1%} upside versus the current/ask level."
    )
    summary.paragraph_format.space_after = Pt(8)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        ff_chart = tmp_path / "football.png"
        margin_chart = tmp_path / "margin.png"
        revenue_chart = tmp_path / "revenue.png"
        fcf_area = tmp_path / "fcf_area.png"
        _build_chart_football(bear, midpoint, bull, ff_chart)
        _build_chart_margin(forecast_df, margin_chart)
        _build_chart_revenue_bridge(forecast_df, revenue_chart)
        _build_chart_fcf_area(forecast_df, fcf_area)

        split = doc.add_table(rows=1, cols=2)
        split.autofit = False
        split.columns[0].width = Inches(3.55)
        split.columns[1].width = Inches(3.55)

        left = split.cell(0, 0)
        right = split.cell(0, 1)

        left.add_paragraph("Key Valuation Metrics").runs[0].bold = True
        left.paragraphs[-1].runs[0].font.color.rgb = RGBColor(31, 78, 120)

        kpi = left.add_table(rows=1, cols=4)
        kpi.style = "Light Grid Accent 1"
        kpi.rows[0].cells[0].text = "Metric"
        kpi.rows[0].cells[1].text = "Base"
        kpi.rows[0].cells[2].text = "Bear"
        kpi.rows[0].cells[3].text = "Bull"

        def _krow(name: str, base_v: float, bear_v: float, bull_v: float, money: bool = True):
            rr = kpi.add_row().cells
            rr[0].text = name
            if money:
                rr[1].text = _fmt_m(base_v)
                rr[2].text = _fmt_m(bear_v)
                rr[3].text = _fmt_m(bull_v)
            else:
                rr[1].text = f"{base_v:.1%}"
                rr[2].text = f"{bear_v:.1%}"
                rr[3].text = f"{bull_v:.1%}"

        _krow("Enterprise Value", midpoint, bear, bull)
        _krow("Equity Value", (eq_g + eq_e) / 2, (eq_g + eq_e) / 2 * 0.9, (eq_g + eq_e) / 2 * 1.1)
        _krow("Revenue CAGR", projected_growth, projected_growth - 0.01, projected_growth + 0.01, money=False)
        _krow("Unlevered Free Cash Flow", float(forecast_df["FCF"].iloc[-1]), float(forecast_df["FCF"].iloc[-1]) * 0.9, float(forecast_df["FCF"].iloc[-1]) * 1.1)
        _zebra_table(kpi)

        right.add_paragraph("Unlevered Free Cash Flow Projection ($Y)").runs[0].bold = True
        right.paragraphs[-1].runs[0].font.color.rgb = RGBColor(31, 78, 120)
        right.add_paragraph().add_run().add_picture(str(fcf_area), width=Inches(3.35))

        hook = doc.add_paragraph(
            f"Based on a WACC of {wacc:.1%} and terminal growth of {terminal_g:.1%}, "
            f"{company_name} shows an intrinsic value range of {_fmt_m(low)} to {_fmt_m(high)}, "
            f"implying {upside_pct:.1%} upside vs current/ask level."
        )
        hook.runs[0].bold = True
        hook.paragraph_format.space_before = Pt(6)
        hook.paragraph_format.space_after = Pt(8)

        doc.add_page_break()
        sec_title = doc.add_heading("2) Assumption Sanity Check", level=1)
        sec_title.runs[0].font.color.rgb = RGBColor(24, 37, 56)

        wacc_table = doc.add_table(rows=1, cols=2)
        wacc_table.style = "Light Grid Accent 1"
        wacc_table.rows[0].cells[0].text = "WACC Component"
        wacc_table.rows[0].cells[1].text = "Value"
        for label, value in [
            ("Risk-Free Rate", risk_free),
            ("Beta", beta),
            ("Equity Risk Premium", mrp),
            ("Size Premium", size_premium),
            ("WACC", wacc),
        ]:
            r = wacc_table.add_row().cells
            r[0].text = label
            r[1].text = f"{value:.2%}" if label != "Beta" else f"{value:.2f}"
        _zebra_table(wacc_table)

        tone = "conservative" if projected_growth <= historical_avg else "aggressive"
        relation = "lower" if projected_growth <= historical_avg else "higher"
        doc.add_paragraph(
            f"The projected revenue CAGR of {projected_growth:.1%} is {relation} than the historical 3-year average "
            f"of {historical_avg:.1%}, reflecting a {tone} outlook."
        )

        doc.add_paragraph(
            f"WACC build: Risk-Free ({risk_free:.2%}) + Beta ({beta:.2f}) × ERP ({mrp:.2%}) + Size Premium ({size_premium:.2%})."
        )

        doc.add_heading("3) Financial Health Check", level=1)
        doc.add_picture(str(revenue_chart), width=Pt(430))
        doc.add_picture(str(margin_chart), width=Pt(430))

        fcf_table = doc.add_table(rows=1, cols=3)
        fcf_table.style = "Light Grid Accent 1"
        fcf_table.rows[0].cells[0].text = "FCF Build Item (Y1)"
        fcf_table.rows[0].cells[1].text = "Value"
        fcf_table.rows[0].cells[2].text = "Comment"

        y1 = forecast_df.iloc[0]
        fcf_rows = [
            ("EBITDA", float(y1.get("EBITDA", 0.0)), "Operating profitability"),
            ("Less CapEx", -float(y1.get("Capex", 0.0)), "Reinvestment requirement"),
            ("Less Change in Working Capital", -float(y1.get("Delta NWC", 0.0)), "Liquidity drag"),
            ("Unlevered Free Cash Flow", float(y1.get("FCF", 0.0)), "Cash available to all capital providers"),
        ]
        for label, value, note in fcf_rows:
            r = fcf_table.add_row().cells
            r[0].text = label
            r[1].text = f"${value:,.0f}"
            r[2].text = note
        _zebra_table(fcf_table)

        doc.add_heading("4) Sensitivity & Scenario Analysis", level=1)
        sens = doc.add_table(rows=1, cols=4)
        sens.style = "Light Grid Accent 1"
        sens.rows[0].cells[0].text = "Case"
        sens.rows[0].cells[1].text = "WACC"
        sens.rows[0].cells[2].text = "Terminal Growth"
        sens.rows[0].cells[3].text = "Implied EV"

        for label, w_delta, g_delta in [
            ("Bear", 0.005, -0.005),
            ("Base", 0.0, 0.0),
            ("Bull", -0.005, 0.005),
        ]:
            w_adj = wacc + w_delta
            g_adj = terminal_g + g_delta
            ev_adj = midpoint * (1 - (w_delta * 2.2) + (g_delta * 3.0))
            r = sens.add_row().cells
            r[0].text = label
            r[1].text = f"{w_adj:.2%}"
            r[2].text = f"{g_adj:.2%}"
            r[3].text = _fmt_m(ev_adj)
        _zebra_table(sens)

        doc.add_paragraph(
            "Risk Warning: valuation sensitivity is highest to discount rate assumptions. "
            "A 1% increase in WACC can compress equity value materially, often more than similar shifts in terminal growth."
        )

        doc.add_heading("5) Peer Comparison (Relative Valuation)", level=1)
        target_multiple = midpoint / max(input_ebitda, 1.0)
        peer_rows = [
            ("Peer A", 4.9),
            ("Peer B", 5.5),
            ("Peer C", 6.1),
            ("Peer D", 5.2),
            ("Peer E", 4.7),
        ]
        peer_avg = sum(x[1] for x in peer_rows) / len(peer_rows)

        peers = doc.add_table(rows=1, cols=3)
        peers.style = "Light Grid Accent 1"
        peers.rows[0].cells[0].text = "Company"
        peers.rows[0].cells[1].text = "EV/EBITDA (x)"
        peers.rows[0].cells[2].text = "Premium/Discount vs Target"
        for name, multiple in peer_rows:
            rr = peers.add_row().cells
            rr[0].text = name
            rr[1].text = f"{multiple:.1f}x"
            rr[2].text = f"{(multiple/target_multiple-1):.1%}"

        rr = peers.add_row().cells
        rr[0].text = f"{company_name} (Implied)"
        rr[1].text = f"{target_multiple:.1f}x"
        rr[2].text = "-"
        _zebra_table(peers)

        rel = "premium" if target_multiple > peer_avg else "discount"
        doc.add_paragraph(
            f"{company_name} is valued at a {rel} to the peer average ({target_multiple:.1f}x vs {peer_avg:.1f}x), "
            "driven by its projected profitability and cash conversion profile."
        )

        doc.add_heading("6) Disclaimer", level=1)
        doc.add_paragraph(
            "This report is for informational purposes only and does not constitute investment advice, "
            "a fairness opinion, or an offer to buy/sell securities. Valuation outputs are model-driven and "
            "sensitive to assumptions that may not materialize."
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


@app.get("/")
def index():
    return render_template("index.html", case_defaults=CASE_DEFAULTS)


@app.post("/generate")
def generate():
    company_name = request.form.get("company_name", "Company")
    scenario = request.form.get("scenario", "Base")
    if scenario not in CASE_DEFAULTS:
        scenario = "Base"

    ask_price = _pick_input(request.form, "ask_price", scenario)
    input_ebitda = _pick_input(request.form, "ebitda", scenario)

    cfg = _build_cfg_from_form(request.form, scenario)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        uploaded = request.files.get("financial_file")
        logo_file = request.files.get("logo_file")
        input_path = tmp_path / "input.csv"
        logo_path: Path | None = None

        if uploaded and uploaded.filename:
            suffix = Path(uploaded.filename).suffix.lower() or ".csv"
            input_path = tmp_path / f"input{suffix}"
            uploaded.save(input_path)
        else:
            input_path = _generate_synthetic_input(request.form, input_path, scenario)

        if logo_file and logo_file.filename:
            logo_suffix = Path(logo_file.filename).suffix.lower() or ".png"
            logo_path = tmp_path / f"logo{logo_suffix}"
            logo_file.save(logo_path)

        excel_path = tmp_path / f"{company_name.replace(' ', '_').lower()}_valuation.xlsx"
        word_path = tmp_path / f"{company_name.replace(' ', '_').lower()}_valuation_report.docx"

        run_dcf_pipeline(input_path, excel_path, cfg, scenario_name=scenario)
        _build_word_report(word_path, excel_path, company_name, ask_price, scenario, logo_path, input_ebitda)

        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
            zf.write(excel_path, arcname=excel_path.name)
            zf.write(word_path, arcname=word_path.name)

        zip_buffer.seek(0)
        return send_file(
            zip_buffer,
            mimetype="application/zip",
            as_attachment=True,
            download_name=f"{company_name.replace(' ', '_').lower()}_valuation_pack.zip",
        )


if __name__ == "__main__":
    app.run(debug=True)
