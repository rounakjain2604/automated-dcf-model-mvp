from __future__ import annotations

import json
from datetime import date
from pathlib import Path
import sys

from docx import Document
from docx.shared import Pt


def load_config_override(path: Path, cfg: DCFConfig) -> DCFConfig:
    payload = json.loads(path.read_text(encoding="utf-8"))

    if "forecast" in payload:
        for key, value in payload["forecast"].items():
            setattr(cfg.forecast, key, value)
    if "wacc" in payload:
        for key, value in payload["wacc"].items():
            setattr(cfg.wacc, key, value)
    if "valuation" in payload:
        for key, value in payload["valuation"].items():
            setattr(cfg.valuation, key, value)

    return cfg


def fmt_m(value: float) -> str:
    return f"${value/1_000_000:,.2f}M"


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    sys.path.insert(0, str(root))

    from src.dcf_generator.config import DCFConfig
    from src.dcf_generator.pipeline import run_dcf_pipeline

    input_path = root / "data" / "business_test_insurance.csv"
    config_path = root / "config.business_wacc24_tg2.json"
    excel_output = root / "output" / "business_test_wacc24_tg2_reportpack.xlsx"
    word_output = root / "output" / "business_test_wacc24_tg2_report.docx"

    cfg = load_config_override(config_path, DCFConfig())
    result = run_dcf_pipeline(input_path, excel_output, cfg, scenario_name="Base")
    valuation = result["valuation_summary"]

    ev_g = float(valuation["Enterprise Value (Gordon)"])
    ev_e = float(valuation["Enterprise Value (Exit)"])
    eq_g = float(valuation["Equity Value (Gordon)"])
    eq_e = float(valuation["Equity Value (Exit)"])

    low = min(ev_g, ev_e)
    high = max(ev_g, ev_e)
    midpoint = (low + high) / 2
    ask_price = 4_000_000
    upside = midpoint - ask_price

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Segoe UI"
    normal_style.font.size = Pt(10.5)

    doc.add_heading("PROJECT [COMPANY NAME] | VALUATION REPORT", level=1)
    doc.add_paragraph(f"Date: {date.today().isoformat()}")
    doc.add_paragraph("Strictly Private & Confidential | Prepared by Rounak Jain, CFA L2 Candidate")

    doc.add_heading("Executive Summary", level=2)
    doc.add_paragraph(f"Implied Enterprise Value Range: {fmt_m(low)} - {fmt_m(high)}")
    doc.add_paragraph(f"Implied Equity Value Range: {fmt_m(min(eq_g, eq_e))} - {fmt_m(max(eq_g, eq_e))}")
    doc.add_paragraph(f"WACC Assumption: 24.0% | Terminal Growth: 2.0% | Revenue Growth: 3.0%")

    doc.add_heading("Investment Narrative", level=2)
    narrative = (
        f"At a conservative 24% discount rate, intrinsic value is approximately {fmt_m(midpoint)}. "
        f"Against a $4.00M ask, buyers are implied to gain about ${upside:,.0f} of immediate equity value."
    )
    doc.add_paragraph(narrative)

    doc.add_heading("Method Detail", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Method"
    hdr[1].text = "Enterprise Value"
    hdr[2].text = "Equity Value"

    row1 = table.add_row().cells
    row1[0].text = "Gordon Growth"
    row1[1].text = fmt_m(ev_g)
    row1[2].text = fmt_m(eq_g)

    row2 = table.add_row().cells
    row2[0].text = "Exit Multiple"
    row2[1].text = fmt_m(ev_e)
    row2[2].text = fmt_m(eq_e)

    doc.add_heading("Files Generated", level=2)
    doc.add_paragraph(f"Excel Model: {excel_output.name}")
    doc.add_paragraph(f"Word Report: {word_output.name}")

    word_output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(word_output)

    print(f"Excel generated: {excel_output}")
    print(f"Word report generated: {word_output}")
    print(f"EV Range: {fmt_m(low)} - {fmt_m(high)}")


if __name__ == "__main__":
    main()
